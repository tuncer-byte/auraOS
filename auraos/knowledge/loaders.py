"""
Document Loaders — dosya tipine göre belge yükleme.

Desteklenen formatlar:
  - .txt, .md — stdlib
  - .csv — stdlib csv
  - .json — stdlib json
  - .pdf — lazy import pdfplumber
  - .docx — lazy import python-docx
  - .html — lazy import beautifulsoup4
"""
from __future__ import annotations
import csv
import io
import json
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional

from auraos.knowledge.document import Document


class DocumentLoader(ABC):
    @abstractmethod
    def load(self, path: str) -> list[Document]:
        ...


class TextLoader(DocumentLoader):
    def load(self, path: str) -> list[Document]:
        text = Path(path).read_text(encoding="utf-8")
        return [Document(content=text, metadata={"source": path, "type": "text"})]


class MarkdownLoader(DocumentLoader):
    def load(self, path: str) -> list[Document]:
        text = Path(path).read_text(encoding="utf-8")
        return [Document(content=text, metadata={"source": path, "type": "markdown"})]


class CSVLoader(DocumentLoader):
    def load(self, path: str) -> list[Document]:
        docs: list[Document] = []
        with open(path, newline="", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for i, row in enumerate(reader):
                content = "\n".join(f"{k}: {v}" for k, v in row.items() if v)
                docs.append(Document(
                    content=content,
                    metadata={"source": path, "type": "csv", "row": i},
                ))
        return docs


class JSONLoader(DocumentLoader):
    def __init__(self, text_key: Optional[str] = None):
        self.text_key = text_key

    def load(self, path: str) -> list[Document]:
        data = json.loads(Path(path).read_text(encoding="utf-8"))
        if isinstance(data, list):
            docs = []
            for i, item in enumerate(data):
                if self.text_key and isinstance(item, dict):
                    content = str(item.get(self.text_key, item))
                else:
                    content = json.dumps(item, ensure_ascii=False, indent=2)
                docs.append(Document(
                    content=content,
                    metadata={"source": path, "type": "json", "index": i},
                ))
            return docs
        content = json.dumps(data, ensure_ascii=False, indent=2)
        return [Document(content=content, metadata={"source": path, "type": "json"})]


class PDFLoader(DocumentLoader):
    def load(self, path: str) -> list[Document]:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("PDF desteği için: pip install pdfplumber")
        docs: list[Document] = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(
                        content=text,
                        metadata={"source": path, "type": "pdf", "page": i + 1},
                    ))
        return docs


class DOCXLoader(DocumentLoader):
    def load(self, path: str) -> list[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("DOCX desteği için: pip install python-docx")
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        return [Document(content=content, metadata={"source": path, "type": "docx"})]


class HTMLLoader(DocumentLoader):
    def load(self, path: str) -> list[Document]:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("HTML desteği için: pip install beautifulsoup4")
        html = Path(path).read_text(encoding="utf-8")
        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        return [Document(content=text, metadata={"source": path, "type": "html"})]


_EXTENSION_MAP: dict[str, type[DocumentLoader]] = {
    ".txt": TextLoader,
    ".md": MarkdownLoader,
    ".csv": CSVLoader,
    ".json": JSONLoader,
    ".pdf": PDFLoader,
    ".docx": DOCXLoader,
    ".html": HTMLLoader,
    ".htm": HTMLLoader,
}


def get_loader(path: str) -> DocumentLoader:
    ext = Path(path).suffix.lower()
    loader_cls = _EXTENSION_MAP.get(ext)
    if loader_cls is None:
        return TextLoader()
    return loader_cls()
